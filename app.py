import streamlit as st
import re
import fitz  # PyMuPDF
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer, util
import google.generativeai as genai
from docx import Document
from docx.shared import RGBColor
import os
from dotenv import load_dotenv
import nltk

# Download the 'punkt' tokenizer data for sentence tokenization
nltk.download('punkt')

# Load environment variables from .env file
load_dotenv()

# Now configure the API key using the environment variable
genai_api_key = os.getenv('GENAI_API_KEY')

if genai_api_key:
    genai.configure(api_key=genai_api_key)
else:
    raise ValueError("API key not found. Please make sure GENAI_API_KEY is set in your .env file.")

# Set up the models
model = genai.GenerativeModel('gemini-1.5-flash')
sentence_model = SentenceTransformer('all-MiniLM-L6-v2')

# Function to load PDF using PyMuPDF (fitz)
def load_pdf_with_fitz(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in doc:
        # Extract text and clean up spacing issues
        page_text = page.get_text("text").replace('-\n', '')  # Handle hyphenation
        text += re.sub(r'\s+', ' ', page_text)  # Normalize multiple spaces to single
    return text.strip()

# Preprocess text by normalizing and tokenizing into sentences
def preprocess_text(text):
    # Additional cleaning to remove unnecessary spaces within words
    text = re.sub(r'(\w)\s+(\w)', r'\1 \2', text)  # Ensures words are properly spaced
    sentences = sent_tokenize(text)  # Sentence tokenization
    return sentences

# Function to extract introduction, clauses, and signatures from the NDA
def extract_intro_clauses_signatures(sentences):
    introduction = []
    clauses = {}
    signatures = []
    current_text = []
    clause_number = 1
    clause_started = False
    signatures_started = False
    current_clause = None

    for sentence in sentences:
        match = re.match(r'^\d+\.\s*', sentence)
        signature_match = re.search(r'(signatures? appear on following page|in witness whereof|by:|name:|title:)', sentence, re.IGNORECASE)

        if not clause_started:
            if match:
                clause_started = True
                if current_text:
                    introduction.extend(current_text)
                current_clause = f"Clause {clause_number}"
                current_text = [sentence.strip()]
            else:
                introduction.append(sentence.strip())
        elif clause_started and not signatures_started:
            if signature_match:
                signatures_started = True
                if current_text:
                    clauses[current_clause] = ' '.join(current_text).strip()
                current_text = [sentence.strip()]
            elif match:
                if current_text:
                    clauses[current_clause] = ' '.join(current_text).strip()
                    clause_number += 1
                    current_clause = f"Clause {clause_number}"
                    current_text = []
                current_text.append(sentence.strip())
            else:
                current_text.append(sentence.strip())
        elif signatures_started:
            signatures.append(sentence.strip())

    if current_text:
        if signatures_started:
            signatures.extend(current_text)
        else:
            clauses[f"Clause {clause_number}"] = ' '.join(current_text).strip()

    return {
        "Introduction": ' '.join(introduction).strip(),
        "Clauses": clauses,
        "Signatures": ' '.join(signatures).strip()
    }

# Use Sentence-BERT to find the best matching clause
def find_best_match(brp_clause, uploaded_clauses):
    best_match = None
    best_similarity = 0
    brp_embedding = sentence_model.encode(brp_clause, convert_to_tensor=True)

    for up_title, up_text in uploaded_clauses.items():
        up_embedding = sentence_model.encode(up_text, convert_to_tensor=True)
        similarity = util.pytorch_cos_sim(brp_embedding, up_embedding).item()

        if similarity > best_similarity:
            best_similarity = similarity
            best_match = (up_title, up_text)

    return best_match, best_similarity

# Updated Gemini prompt with focus on essential differences
def analyze_with_gemini(segment1, segment2):
    prompt = (
        f"Identify the essential differences between the following two NDA segments. Focus only on the differences that could have legal implications, ignoring formatting, typos, or stylistic changes. "
        f"Summarize the key points of deviation briefly."
        f"\n\nSegment 1 (BRP NDA): {segment1}\n\n"
        f"Segment 2 ({uploaded_file_name}): {segment2}\n\n"
    )
    
    try:
        # Gemini API call to generate text
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error in processing the segments with Gemini: {e}"

# Refine the output to focus on key differences
def refine_deviations_with_gemini(deviation):
    prompt = (
        f"Highlight only the essential differences that could have legal implications. Keep the summary as concise as possible:"
        f"\n\n{deviation}"
    )
    
    try:
        # Gemini API call to refine the output
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error in refining the deviations with Gemini: {e}"

# Enhanced comparison to avoid false positives and refine deviations
def compare_documents(brp_clauses, uploaded_clauses, similarity_threshold=0.75):  # Adjusted threshold for more sensitivity
    comparisons = []
    brp_matched_clauses = set()
    uploaded_matched_clauses = set()
    
    for brp_title, brp_segment in brp_clauses.items():
        best_match, best_similarity = find_best_match(brp_segment, uploaded_clauses)

        if best_match and best_similarity >= similarity_threshold:
            analysis = analyze_with_gemini(brp_segment, best_match[1])

            # Further refine the deviations
            refined_deviation = refine_deviations_with_gemini(analysis)

            if brp_segment.strip() != best_match[1].strip():
                comparisons.append({
                    'BRP Clause Title': brp_title,
                    'Uploaded NDA Match Title': best_match[0],
                    'Deviation': refined_deviation
                })
            else:
                comparisons.append({
                    'BRP Clause Title': brp_title,
                    'Uploaded NDA Match Title': best_match[0],
                    'Deviation': 'No significant deviations found.'
                })
            brp_matched_clauses.add(brp_title)
            uploaded_matched_clauses.add(best_match[0])
        else:
            comparisons.append({
                'BRP Clause Title': brp_title,
                'Uploaded NDA Match Title': None,
                'Deviation': f'This clause is missing in the {uploaded_file_name} and should be reviewed or added.'
            })

    # Detect new clauses in Uploaded NDA that are not in BRP NDA
    for up_title in uploaded_clauses.keys():
        if up_title not in uploaded_matched_clauses:
            comparisons.append({
                'BRP Clause Title': None,
                'Uploaded NDA Match Title': up_title,
                'Deviation': f'This is a new clause in the {uploaded_file_name} that is not present in the BRP NDA and should be reviewed.'
            })

    return comparisons

# Function to generate a DOCX file with comparison results
def generate_highlighted_docx(comparisons, reviewed_items, file_name="nda_comparison_results.docx"):
    doc = Document()
    doc.add_heading('NDA Clause Comparison Results', 0)
    
    suggestions = []  # To store the summary bullet points
    review_needed = []  # To store clauses that need further human review
    
    for idx, comparison in enumerate(comparisons):
        doc.add_heading(comparison['BRP Clause Title'] or "New Clause", level=2)
        doc.add_paragraph(f"Uploaded NDA Match Title: {comparison['Uploaded NDA Match Title']}")
        deviation_paragraph = doc.add_paragraph("Deviation: ")
        deviation_run = deviation_paragraph.add_run(comparison['Deviation'])
        deviation_run.font.bold = True  # Highlight deviations in bold
        deviation_run.font.color.rgb = RGBColor(255, 0, 0)  # Highlight deviations in red
        
        # If the user marked this clause for further review, note it in the document
        if reviewed_items[idx]:
            review_needed.append(comparison['BRP Clause Title'] or "New Clause")
            doc.add_paragraph("**Marked for further human review**", style='Intense Quote')
        
        # Add the key suggestion to the summary list
        if "suggest specific changes" in comparison['Deviation'].lower():
            suggestions.append(f"Modify {comparison['Uploaded NDA Match Title']} to align with the BRP NDA as per the highlighted differences.")
            
        doc.add_paragraph()  # Add space before the next item
    
    # Add a summary section with bullet points
    if suggestions:
        doc.add_heading('Summary of Key Suggestions', level=1)
        for suggestion in suggestions:
            doc.add_paragraph(suggestion, style='List Bullet')
    
    # Add a section listing clauses that require further human review
    if review_needed:
        doc.add_heading('Clauses Marked for Further Human Review', level=1)
        for clause in review_needed:
            doc.add_paragraph(clause, style='List Bullet')
    
    doc.save(file_name)
    return file_name

# Streamlit UI
st.title("NDA Clause Comparison Tool")

# File upload for BRP NDA Template and Uploaded NDA
brp_file = st.file_uploader("Upload BRP NDA Template (PDF)", type="pdf")
uploaded_file = st.file_uploader("Upload NDA for Comparison (PDF)", type="pdf")

if brp_file and uploaded_file:
    with st.spinner('Processing the NDAs...'):
        # Get the uploaded file name
        uploaded_file_name = os.path.basename(uploaded_file.name)

        # Load and preprocess the texts from both PDFs
        brp_nda_text = load_pdf_with_fitz(brp_file)
        uploaded_nda_text = load_pdf_with_fitz(uploaded_file)

        # Display the extracted texts for review
        st.subheader("Extracted Text from BRP NDA Template")
        st.text_area("BRP NDA Text", brp_nda_text, height=200)

        st.subheader(f"Extracted Text from {uploaded_file_name}")
        st.text_area(f"{uploaded_file_name} Text", uploaded_nda_text, height=200)

        brp_nda_sentences = preprocess_text(brp_nda_text)
        uploaded_nda_sentences = preprocess_text(uploaded_nda_text)

        # Extract the sections from both NDAs
        brp_sections = extract_intro_clauses_signatures(brp_nda_sentences)
        uploaded_sections = extract_intro_clauses_signatures(uploaded_nda_sentences)

        st.write("Clauses extracted. Comparing...")

        # Compare the clauses using Gemini
        gemini_comparisons = compare_documents(brp_sections["Clauses"], uploaded_sections["Clauses"])

        st.write("Comparison Results:")

        reviewed_items = []  # Store whether each clause needs further review

        # Display the checklist directly in the app
        st.subheader("Checklist of Items to Review/Modify")
        for idx, comparison in enumerate(gemini_comparisons):
            st.markdown(f"### {comparison['BRP Clause Title'] or 'New Clause'}")
            st.write(f"**{uploaded_file_name} Match Title:** {comparison['Uploaded NDA Match Title'] if comparison['Uploaded NDA Match Title'] else 'N/A'}")
            st.write(f"**Deviation:** {comparison['Deviation']}")
            needs_review = st.checkbox(f"Mark {comparison['BRP Clause Title'] or 'New Clause'} for further human review", key=f"review_{idx}")
            reviewed_items.append(needs_review)  # Track the checkbox status

            st.write("---")

        st.success("Comparison complete!")

        # Generate and provide a DOCX file for download with highlighted deviations and checklist
        comparison_file_name = generate_highlighted_docx(gemini_comparisons, reviewed_items)
        st.download_button(
            label="Download Comparison Results",
            data=open(comparison_file_name, "rb").read(),
            file_name=comparison_file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Dynamically set the port for Cloud Run
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))  # Use port 8501 locally, or 8080 in Cloud Run
    st.run("app.py", f"--server.port={port}")

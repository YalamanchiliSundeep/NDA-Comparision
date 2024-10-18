import streamlit as st
import re
import fitz  # PyMuPDF
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import RGBColor
import os
from dotenv import load_dotenv
import nltk
import openai

# Download the 'punkt' tokenizer data for sentence tokenization
nltk.download('punkt_tab')

# Load environment variables from .env file
load_dotenv()

# OpenAI API key setup
openai.api_key = os.getenv('OPENAI_API_KEY')

# Set up the models
sentence_model = SentenceTransformer('all-MiniLM-L6-v2')

# Function to load PDF using PyMuPDF (fitz)
def load_pdf_with_fitz(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text("text").replace('-\n', '')  # Handle hyphenation
        text += re.sub(r'\s+', ' ', page_text)  # Normalize multiple spaces to single
    return text.strip()

# Preprocess text by normalizing and tokenizing into sentences
def preprocess_text(text):
    text = re.sub(r'(\w)\s+(\w)', r'\1 \2', text)  # Ensures words are properly spaced
    sentences = sent_tokenize(text)  # Sentence tokenization
    return sentences

# Function to extract introduction, clauses, and signatures from the NDA
def extract_intro_clauses_signatures(sentences):
    introduction = []
    clauses = {}
    signatures = []
    current_text = []
    clause_number = 1
    clause_started = False
    signatures_started = False
    current_clause = None

    for sentence in sentences:
        match = re.match(r'^\d+\.\s*', sentence)
        signature_match = re.search(r'(signatures? appear on following page|in witness whereof|by:|name:|title:)', sentence, re.IGNORECASE)

        if not clause_started:
            if match:
                clause_started = True
                if current_text:
                    introduction.extend(current_text)
                current_clause = f"Clause {clause_number}"
                current_text = [sentence.strip()]
            else:
                introduction.append(sentence.strip())
        elif clause_started and not signatures_started:
            if signature_match:
                signatures_started = True
                if current_text:
                    clauses[current_clause] = ' '.join(current_text).strip()
                current_text = [sentence.strip()]
            elif match:
                if current_text:
                    clauses[current_clause] = ' '.join(current_text).strip()
                    clause_number += 1
                    current_clause = f"Clause {clause_number}"
                    current_text = []
                current_text.append(sentence.strip())
            else:
                current_text.append(sentence.strip())
        elif signatures_started:
            signatures.append(sentence.strip())

    if current_text:
        if signatures_started:
            signatures.extend(current_text)
        else:
            clauses[f"Clause {clause_number}"] = ' '.join(current_text).strip()

    return {
        "Introduction": ' '.join(introduction).strip(),
        "Clauses": clauses,
        "Signatures": ' '.join(signatures).strip()
    }

# Use Sentence-BERT to find the best matching clause
def find_best_match(brp_clause, uploaded_clauses):
    best_match = None
    best_similarity = 0
    brp_embedding = sentence_model.encode(brp_clause, convert_to_tensor=True)

    for up_title, up_text in uploaded_clauses.items():
        up_embedding = sentence_model.encode(up_text, convert_to_tensor=True)
        similarity = util.pytorch_cos_sim(brp_embedding, up_embedding).item()

        if similarity > best_similarity:
            best_similarity = similarity
            best_match = (up_title, up_text)

    return best_match, best_similarity

# Updated API: Analyze deviations with GPT
def analyze_with_gpt(segment1, segment2):
    messages = [
        {"role": "system", "content": "You are a legal assistant tasked with comparing NDA clauses and identifying key legal differences."},
        {"role": "user", "content": (
            f"Compare the following two NDA clauses and identify key legal differences. "
            f"Focus only on differences that have legal implications, and ignore formatting or stylistic changes."
            f"\n\nClause 1: {segment1}\n\nClause 2: {segment2}"
        )}
    ]
    
    try:
        response = openai.completions.create(  # new method in API >= 1.0.0
            model="gpt-3.5-turbo",  # Or "gpt-4"
            messages=messages,
            max_tokens=500,  # Adjust token limit based on expected response size
            temperature=0.5
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        return f"Error in processing the segments with GPT: {e}"

# Updated API: Refine deviations with GPT
def refine_deviation_with_gpt(deviation):
    messages = [
        {"role": "system", "content": "You are a legal assistant tasked with summarizing key legal differences concisely."},
        {"role": "user", "content": (
            f"Summarize the following legal differences in the most concise way possible, focusing only on essential legal implications."
            f"\n\n{deviation}"
        )}
    ]
    
    try:
        response = openai.completions.create(  # new method in API >= 1.0.0
            model="gpt-3.5-turbo",  # Or "gpt-4"
            messages=messages,
            max_tokens=300,  # Adjust token limit for brevity
            temperature=0.3
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        return f"Error in refining the deviations with GPT: {e}"

# Comparison function with analysis and refinement
def compare_documents(brp_clauses, uploaded_clauses, similarity_threshold=0.75):
    comparisons = []
    brp_matched_clauses = set()
    uploaded_matched_clauses = set()
    
    for brp_title, brp_segment in brp_clauses.items():
        best_match, best_similarity = find_best_match(brp_segment, uploaded_clauses)

        if best_match and best_similarity >= similarity_threshold:
            # Step 1: Analyze deviations
            deviation = analyze_with_gpt(brp_segment, best_match[1])

            # Step 2: Refine the deviations
            refined_deviation = refine_deviation_with_gpt(deviation)

            if brp_segment.strip() != best_match[1].strip():
                comparisons.append({
                    'BRP Clause Title': brp_title,
                    'Uploaded NDA Match Title': best_match[0],
                    'Deviation': refined_deviation
                })
            else:
                comparisons.append({
                    'BRP Clause Title': brp_title,
                    'Uploaded NDA Match Title': best_match[0],
                    'Deviation': 'No significant deviations found.'
                })
            brp_matched_clauses.add(brp_title)
            uploaded_matched_clauses.add(best_match[0])
        else:
            comparisons.append({
                'BRP Clause Title': brp_title,
                'Uploaded NDA Match Title': None,
                'Deviation': f'This clause is missing in the uploaded NDA and should be reviewed or added.'
            })

    # Detect new clauses in Uploaded NDA that are not in BRP NDA
    for up_title in uploaded_clauses.keys():
        if up_title not in uploaded_matched_clauses:
            comparisons.append({
                'BRP Clause Title': None,
                'Uploaded NDA Match Title': up_title,
                'Deviation': f'This is a new clause in the uploaded NDA that is not present in the BRP NDA and should be reviewed.'
            })

    return comparisons

# Function to generate a DOCX file with comparison results
def generate_highlighted_docx(comparisons, reviewed_items, file_name="nda_comparison_results.docx"):
    doc = Document()
    doc.add_heading('NDA Clause Comparison Results', 0)
    
    suggestions = []  # To store the summary bullet points
    review_needed = []  # To store clauses that need further human review
    
    for idx, comparison in enumerate(comparisons):
        doc.add_heading(comparison['BRP Clause Title'] or "New Clause", level=2)
        doc.add_paragraph(f"Uploaded NDA Match Title: {comparison['Uploaded NDA Match Title']}")
        deviation_paragraph = doc.add_paragraph("Deviation: ")
        deviation_run = deviation_paragraph.add_run(comparison['Deviation'])
        deviation_run.font.bold = True  # Highlight deviations in bold
        deviation_run.font.color.rgb = RGBColor(255, 0, 0)  # Highlight deviations in red
        
        # If the user marked this clause for further review, note it in the document
        if reviewed_items[idx]:
            review_needed.append(comparison['BRP Clause Title'] or "New Clause")
            doc.add_paragraph("**Marked for further human review**", style='Intense Quote')
        
        # Add the key suggestion to the summary list
        if "suggest specific changes" in comparison['Deviation'].lower():
            suggestions.append(f"Modify {comparison['Uploaded NDA Match Title']} to align with the BRP NDA as per the highlighted differences.")
            
        doc.add_paragraph()  # Add space before the next item
    
    # Add a summary section with bullet points
    if suggestions:
        doc.add_heading('Summary of Key Suggestions', level=1)
        for suggestion in suggestions:
            doc.add_paragraph(suggestion, style='List Bullet')
    
    # Add a section listing clauses that require further human review
    if review_needed:
        doc.add_heading('Clauses Marked for Further Human Review', level=1)
        for clause in review_needed:
            doc.add_paragraph(clause, style='List Bullet')
    
    doc.save(file_name)
    return file_name

# Streamlit UI
st.title("NDA Clause Comparison Tool")

# File upload for BRP NDA Template and Uploaded NDA
brp_file = st.file_uploader("Upload BRP NDA Template (PDF)", type="pdf")
uploaded_file = st.file_uploader("Upload NDA for Comparison (PDF)", type="pdf")

if brp_file and uploaded_file:
    with st.spinner('Processing the NDAs...'):
        # Load and preprocess the texts from both PDFs
        brp_nda_text = load_pdf_with_fitz(brp_file)
        uploaded_nda_text = load_pdf_with_fitz(uploaded_file)

        st.subheader("Extracted Text from BRP NDA Template")
        st.text_area("BRP NDA Text", brp_nda_text, height=200)

        st.subheader(f"Extracted Text from Uploaded NDA")
        st.text_area(f"Uploaded NDA Text", uploaded_nda_text, height=200)

        brp_nda_sentences = preprocess_text(brp_nda_text)
        uploaded_nda_sentences = preprocess_text(uploaded_nda_text)

        brp_sections = extract_intro_clauses_signatures(brp_nda_sentences)
        uploaded_sections = extract_intro_clauses_signatures(uploaded_nda_sentences)

        st.write("Clauses extracted. Comparing...")

        # Compare the clauses using GPT and refine deviations
        gpt_comparisons = compare_documents(brp_sections["Clauses"], uploaded_sections["Clauses"])

        st.write("Comparison Results:")

        reviewed_items = []  # Store whether each clause needs further review

        # Display the checklist directly in the app
        st.subheader("Checklist of Items to Review/Modify")
        for idx, comparison in enumerate(gpt_comparisons):
            st.markdown(f"### {comparison['BRP Clause Title'] or 'New Clause'}")
            st.write(f"**Uploaded NDA Match Title:** {comparison['Uploaded NDA Match Title'] if comparison['Uploaded NDA Match Title'] else 'N/A'}")
            st.write(f"**Deviation:** {comparison['Deviation']}")
            needs_review = st.checkbox(f"Mark {comparison['BRP Clause Title'] or 'New Clause'} for further human review", key=f"review_{idx}")
            reviewed_items.append(needs_review)  # Track the checkbox status

            st.write("---")

        st.success("Comparison complete!")

        # Generate and provide a DOCX file for download with highlighted deviations and checklist
        comparison_file_name = generate_highlighted_docx(gpt_comparisons, reviewed_items)
        st.download_button(
            label="Download Comparison Results",
            data=open(comparison_file_name, "rb").read(),
            file_name=comparison_file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )